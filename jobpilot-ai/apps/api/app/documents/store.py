"""Persistence, serialization, and export for generated documents."""

from __future__ import annotations

import json
import os
from pathlib import Path
from typing import Any

from docx import Document
from sqlalchemy.orm import Session

from app.core.audit import record_audit
from app.models.entities import DocumentFormat, DocumentType, GeneratedDocument, JobPosting
from app.services.documents import profile_payload, public_dict


def persist_document(
    db: Session,
    user_id: int,
    job: JobPosting,
    document_type: DocumentType,
    *,
    title: str,
    content: dict[str, Any],
    markdown: str,
    plain_text: str = "",
    quality: dict[str, Any] | None = None,
    model_used: str,
) -> GeneratedDocument:
    record = GeneratedDocument(
        user_id=user_id,
        job_id=job.id,
        type=document_type,
        format=DocumentFormat.json,
        title=title,
        content=_json_safe(content),
        content_markdown=markdown,
        plain_text=plain_text,
        quality=_json_safe(quality or {}),
        source_profile_snapshot=_json_safe(profile_payload(db, user_id)),
        job_snapshot=_json_safe(public_dict(job)),
        model_used=model_used,
        format_version="v2",
    )
    db.add(record)
    record_audit(db, user_id, "document_generated", {"job_id": job.id, "type": document_type.value})
    db.commit()
    db.refresh(record)
    return record


def serialize_document(
    record: GeneratedDocument,
    *,
    warnings: list[str] | None = None,
    unsupported_claims_removed: list[str] | None = None,
) -> dict[str, Any]:
    quality = record.quality or {}
    return {
        "document_id": record.id,
        "document_type": record.type.value,
        "title": record.title,
        "content": record.content,
        "markdown": record.content_markdown or "",
        "plain_text": record.plain_text or "",
        "quality": quality,
        "model_used": record.model_used,
        # Kept for backward compatibility; the frontend now reads quality.warnings.
        "warnings": warnings if warnings is not None else quality.get("warnings", []),
        "unsupported_claims_removed": (
            unsupported_claims_removed
            if unsupported_claims_removed is not None
            else quality.get("unsupported_claims_removed", [])
        ),
        "download_urls": {
            "docx": f"/jobs/documents/{record.id}/download/docx",
            "pdf": f"/jobs/documents/{record.id}/download/pdf",
        },
    }


def export_document(record: GeneratedDocument, fmt: DocumentFormat) -> str:
    out_dir = Path(os.getenv("UPLOAD_DIR", "uploads")).parent / "generated"
    out_dir.mkdir(parents=True, exist_ok=True)
    path = out_dir / f"document-{record.id}.{fmt.value}"
    content = record.content or {}
    is_resume = record.type == DocumentType.resume
    if fmt == DocumentFormat.docx:
        _render_docx(content, is_resume, path)
        record.docx_file_path = str(path)
    elif fmt == DocumentFormat.pdf:
        _render_pdf(record.plain_text or record.content_markdown or "", is_resume, path)
        record.pdf_file_path = str(path)
    else:
        path.write_text(record.plain_text or "", encoding="utf-8")
    return str(path)


# --------------------------------------------------------------------------- #
# DOCX rendering from structured content (ATS-friendly, matches the preview)
# --------------------------------------------------------------------------- #
def _render_docx(content: dict[str, Any], is_resume: bool, path: Path) -> None:
    doc = Document()
    if is_resume:
        _render_resume_docx(doc, content)
    else:
        _render_cover_docx(doc, content)
    doc.save(path)


def _render_resume_docx(doc, content: dict[str, Any]) -> None:
    header = content.get("header") or {}
    if header.get("full_name"):
        doc.add_heading(header["full_name"], level=0)
    contact = " | ".join(
        str(v) for v in [header.get("email"), header.get("phone"), header.get("location"), *(header.get("links") or [])] if v
    )
    if contact:
        doc.add_paragraph(contact)

    if content.get("summary"):
        doc.add_heading("Professional Summary", level=1)
        doc.add_paragraph(content["summary"])

    skills = [g for g in content.get("skills") or [] if g.get("items")]
    if skills:
        doc.add_heading("Core Skills", level=1)
        for group in skills:
            prefix = f"{group['category']}: " if group.get("category") else ""
            doc.add_paragraph(prefix + ", ".join(group["items"]))

    if content.get("experience"):
        doc.add_heading("Professional Experience", level=1)
        for exp in content["experience"]:
            head = doc.add_paragraph()
            head.add_run(" — ".join(filter(None, [exp.get("title"), exp.get("company")]))).bold = True
            meta = " | ".join(filter(None, [exp.get("location"), exp.get("dates")]))
            if meta:
                doc.add_paragraph(meta)
            for bullet in exp.get("bullets") or []:
                doc.add_paragraph(bullet, style="List Bullet")

    if content.get("projects"):
        doc.add_heading("Selected Projects", level=1)
        for proj in content["projects"]:
            head = doc.add_paragraph()
            head.add_run(proj.get("name", "")).bold = True
            tech = ", ".join(proj.get("technologies") or [])
            if tech:
                doc.add_paragraph(tech)
            for bullet in proj.get("bullets") or []:
                doc.add_paragraph(bullet, style="List Bullet")

    if content.get("education"):
        doc.add_heading("Education", level=1)
        for edu in content["education"]:
            doc.add_paragraph(
                ", ".join(filter(None, [edu.get("school"), edu.get("degree"), edu.get("dates"), edu.get("details")]))
            )

    awards = (content.get("awards") or []) + (content.get("certifications") or [])
    awards = [a for a in awards if (a.get("name") if isinstance(a, dict) else a)]
    if awards:
        doc.add_heading("Awards & Certifications", level=1)
        for award in awards:
            name = award.get("name") if isinstance(award, dict) else str(award)
            doc.add_paragraph(name, style="List Bullet")


def _render_cover_docx(doc, content: dict[str, Any]) -> None:
    doc.add_paragraph(content.get("date", ""))
    doc.add_paragraph(content.get("recipient", "Hiring Team"))
    if content.get("company"):
        doc.add_paragraph(content["company"])
    doc.add_paragraph("")
    doc.add_paragraph(content.get("greeting", "Dear Hiring Team,"))
    for para in content.get("paragraphs") or []:
        doc.add_paragraph(para)
    doc.add_paragraph("")
    doc.add_paragraph(content.get("closing", "Best regards,"))
    doc.add_paragraph(content.get("signature", ""))


def _render_pdf(text: str, is_resume: bool, path: Path) -> None:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    pdf = canvas.Canvas(str(path), pagesize=letter)
    _width, height = letter
    y = height - 72
    for raw in text.splitlines():
        line = raw.rstrip()
        stripped = line.strip()
        # Uppercase section headers (from plain_text) render bold.
        is_heading = bool(stripped) and stripped == stripped.upper() and len(stripped) < 40 and not stripped.startswith("•")
        font, size = ("Helvetica-Bold", 12) if is_heading else ("Helvetica", 10)
        pdf.setFont(font, size)
        if not stripped:
            y -= 8
            continue
        for chunk in _wrap(line, 95):
            pdf.drawString(72, y, chunk)
            y -= size + 4
            if y < 72:
                pdf.showPage()
                y = height - 72
    pdf.save()


def _json_safe(value: Any) -> Any:
    """Recursively convert datetimes/dates to strings so snapshots persist as JSON."""
    return json.loads(json.dumps(value, default=str))


def _wrap(text: str, width: int) -> list[str]:
    words, lines, current = text.split(), [], ""
    for word in words:
        if len(current) + len(word) + 1 > width:
            lines.append(current)
            current = word
        else:
            current = f"{current} {word}".strip()
    if current:
        lines.append(current)
    return lines or [""]
