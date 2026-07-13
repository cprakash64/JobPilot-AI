"""Tests for structured resume/cover-letter content, quality, and clean export."""

import io
from collections.abc import Generator
from datetime import UTC, datetime, timedelta

import pytest
from docx import Document
from fastapi.testclient import TestClient
from sqlalchemy import create_engine
from sqlalchemy.orm import Session, sessionmaker
from sqlalchemy.pool import StaticPool

from app.api.deps import get_db
from app.db.base import Base
from app.main import app
from app.models import entities as E
from app.models import entities  # noqa: F401


@pytest.fixture()
def client() -> Generator[TestClient, None, None]:
    engine = create_engine("sqlite://", connect_args={"check_same_thread": False}, poolclass=StaticPool)
    TestingSessionLocal = sessionmaker(bind=engine, autoflush=False, autocommit=False)
    Base.metadata.create_all(bind=engine)

    def override_get_db() -> Generator[Session, None, None]:
        db = TestingSessionLocal()
        try:
            yield db
        finally:
            db.close()

    app.dependency_overrides[get_db] = override_get_db
    try:
        yield TestClient(app)
    finally:
        app.dependency_overrides.clear()
        Base.metadata.drop_all(bind=engine)
        engine.dispose()


def setup(client: TestClient) -> tuple[dict, int]:
    token = client.post("/auth/signup", json={"email": "struct@example.com", "password": "password123"}).json()
    headers = {"Authorization": f"Bearer {token['access_token']}"}
    client.put("/profile", headers=headers, json={
        "full_name": "Chandra Pandey", "phone": "602-816-1309", "location_city": "Phoenix", "location_state": "AZ",
        "target_roles": ["Backend Engineer"], "target_levels": ["Junior"], "preferred_locations": ["United States"],
        "remote_preference": "everything", "skills": ["Python", "FastAPI", "PostgreSQL"],
        "linkedin_url": "https://linkedin.com/in/cp",
    })
    client.put("/profile/career", headers=headers, json={
        "education": [{"school": "Arizona State University", "degree": "BS", "major": "Computer Science"}],
        "experience": [{"company": "Cardinal Health", "title": "ML Engineer Intern",
                        "bullets": ["Built Python services", "Shipped RAG search"], "technologies": ["Python", "FastAPI"]}],
        "projects": [{"name": "Luna AI", "description": "Video platform", "bullets": ["Built pipeline"], "technologies": ["Python"]}],
        "certifications": [], "awards": [],
    })
    db = next(app.dependency_overrides[get_db]())
    src = E.JobSource(name="Acme", type="greenhouse", base_url="x", enabled=True, supports_api=True)
    db.add(src)
    db.flush()
    job = E.JobPosting(
        source_id=src.id, external_id="1", title="Backend Engineer", company="Acme",
        location="Remote, United States", remote_type="remote", posted_at=datetime.now(UTC) - timedelta(days=1),
        discovered_at=datetime.now(UTC), application_url="https://boards.greenhouse.io/acme/1",
        source_url="https://boards.greenhouse.io/acme/1", description_raw="",
        description_clean="Requirements: Python, FastAPI, Rust, Go.", required_skills=["Python", "FastAPI", "Rust", "Go"],
        hash_for_deduplication="h1",
    )
    db.add(job)
    db.commit()
    job_id = job.id
    db.close()
    return headers, job_id


# --------------------------------------------------------------------------- #
# Resume structure + quality
# --------------------------------------------------------------------------- #
def test_resume_returns_structured_content(client: TestClient) -> None:
    headers, job_id = setup(client)
    body = client.post(f"/jobs/{job_id}/generate-resume", headers=headers).json()
    content = body["content"]
    assert content["header"]["full_name"] == "Chandra Pandey"
    assert content["header"]["location"] == "Phoenix, AZ"
    assert "https://linkedin.com/in/cp" in content["header"]["links"]
    # Skills are grouped into real categories {category, items} (not one comma dump).
    categories = {group["category"] for group in content["skills"]}
    assert categories and "Core Skills" not in categories
    all_items = {item for group in content["skills"] for item in group["items"]}
    assert "Python" in all_items
    # Experience carries title/company/dates/bullets.
    exp = content["experience"][0]
    assert exp["company"] == "Cardinal Health" and exp["title"] == "ML Engineer Intern"
    assert set(exp) >= {"title", "company", "location", "dates", "bullets"}


def test_resume_quality_object(client: TestClient) -> None:
    headers, job_id = setup(client)
    body = client.post(f"/jobs/{job_id}/generate-resume", headers=headers).json()
    quality = body["quality"]
    assert quality["ats_friendly"] is True
    assert quality["job_tailored"] is True
    assert quality["estimated_page_count"] == 1  # junior single-page
    # Rust/Go are required by the JD but not in the profile -> not claimed.
    assert "Rust" in quality["missing_job_skills_not_claimed"]
    assert "Go" in quality["missing_job_skills_not_claimed"]
    resume_skills = {i for g in body["content"]["skills"] for i in g["items"]}
    assert "Rust" not in resume_skills and "Go" not in resume_skills


def test_resume_plain_text_has_no_markdown(client: TestClient) -> None:
    headers, job_id = setup(client)
    body = client.post(f"/jobs/{job_id}/generate-resume", headers=headers).json()
    plain = body["plain_text"]
    assert plain
    assert "#" not in plain
    assert "**" not in plain
    assert "Chandra Pandey" in plain


def test_resume_docx_export_is_clean(client: TestClient) -> None:
    headers, job_id = setup(client)
    doc_id = client.post(f"/jobs/{job_id}/generate-resume", headers=headers).json()["document_id"]
    response = client.get(f"/jobs/documents/{doc_id}/download/docx", headers=headers)
    assert response.status_code == 200
    assert response.headers["content-type"] == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    document = Document(io.BytesIO(response.content))
    text = "\n".join(p.text for p in document.paragraphs)
    assert "Chandra Pandey" in text
    assert "#" not in text and "**" not in text
    # Standard ATS section headings present.
    assert "Professional Experience" in text


# --------------------------------------------------------------------------- #
# Cover letter structure
# --------------------------------------------------------------------------- #
def test_cover_letter_structured_paragraphs(client: TestClient) -> None:
    headers, job_id = setup(client)
    body = client.post(f"/jobs/{job_id}/generate-cover-letter", headers=headers).json()
    content = body["content"]
    assert set(content) >= {"date", "recipient", "company", "role", "greeting", "paragraphs", "closing", "signature"}
    assert content["company"] == "Acme"
    assert content["role"] == "Backend Engineer"
    assert content["signature"] == "Chandra Pandey"
    assert len(content["paragraphs"]) == 3
    joined = " ".join(content["paragraphs"])
    assert "Acme" in joined and "Backend Engineer" in joined
    assert 150 <= len(joined.split()) <= 300


def test_cover_letter_does_not_claim_missing_skill(client: TestClient) -> None:
    headers, job_id = setup(client)
    body = client.post(f"/jobs/{job_id}/generate-cover-letter", headers=headers).json()
    joined = " ".join(body["content"]["paragraphs"]).lower()
    # The JD requires Rust; the profile does not have it, so it must not be claimed.
    assert "rust" not in joined


def test_cover_letter_docx_is_a_letter(client: TestClient) -> None:
    headers, job_id = setup(client)
    doc_id = client.post(f"/jobs/{job_id}/generate-cover-letter", headers=headers).json()["document_id"]
    response = client.get(f"/jobs/documents/{doc_id}/download/docx", headers=headers)
    assert response.status_code == 200
    text = "\n".join(p.text for p in Document(io.BytesIO(response.content)).paragraphs)
    assert "Dear Hiring Team," in text
    assert "Acme" in text
    assert "#" not in text


# --------------------------------------------------------------------------- #
# Fallback still structured
# --------------------------------------------------------------------------- #
def test_ai_unavailable_still_returns_structured_content(client: TestClient) -> None:
    from app.ai.provider import ai_provider
    assert ai_provider.client is None  # no key in tests -> template mode
    headers, job_id = setup(client)
    resume = client.post(f"/jobs/{job_id}/generate-resume", headers=headers).json()
    assert resume["content"]["skills"][0]["items"]  # grouped skills present
    assert resume["quality"]["ats_friendly"] is True
    cover = client.post(f"/jobs/{job_id}/generate-cover-letter", headers=headers).json()
    assert len(cover["content"]["paragraphs"]) == 3
    assert any("template mode" in w for w in resume["warnings"])
